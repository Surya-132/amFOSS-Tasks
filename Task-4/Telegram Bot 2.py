import requests
import csv
import io
from typing import Final
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler, CallbackQueryHandler
from docx import Document
from io import BytesIO
BOT_TOKEN = '6921875478:AAHjPqIVJrCqLt8j-206MTSF_ZIAMJ4PhjY'
API_KEY = 'AIzaSyDxhTNEXzvWblGDS0vuPLMayFG7bpyMTjk'
BOT_USERNAME : Final ='@Why_So_Long_Bot'
ASK_GENRE = 3
ASK_BOOK_NAME = 4
READING_LIST_FILE = "reading_list.docx"

def initialize_reading_list():
    doc = Document()
    doc.add_heading('Reading List', 0)
    doc.save(READING_LIST_FILE)

def add_book(title, link):
    doc = Document(READING_LIST_FILE)
    doc.add_paragraph(f"{title} - {link}")
    doc.save(READING_LIST_FILE)

def delete_book(title):
    doc = Document(READING_LIST_FILE)
    paragraphs = doc.paragraphs
    doc._body.clear_content()
    doc.add_heading('Reading List', 0)

    for paragraph in paragraphs:
        if title not in paragraph.text:
            doc.add_paragraph(paragraph.text)
    
    doc.save(READING_LIST_FILE)

async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    if query.data == 'add_book':
        await query.edit_message_text(text="Please send the book title and preview link separated by a comma.")
        context.user_data['action'] = 'add_book'
    elif query.data == 'delete_book':
        await query.edit_message_text(text="Please send the book title to delete:")
        context.user_data['action'] = 'delete_book'
    elif query.data == 'view_list':
        await query.message.reply_document(document=open(READING_LIST_FILE, 'rb'))
      
async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_action = context.user_data.get('action')

    if user_action == 'add_book':
        try:
            title, link = update.message.text.split(',')
            add_book(title.strip(), link.strip())
            await update.message.reply_text("Book added successfully!")
        except ValueError:
            await update.message.reply_text("Please send the book title and preview link in the correct format.")
    elif user_action == 'delete_book':
        title = update.message.text.strip()
        delete_book(title)
        await update.message.reply_text("Book deleted successfully!")

    
    context.user_data['action'] = None

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(" Welcome to Why_So_Long! Looking for your next great read? You've found the perfect place. Happy reading!")

async def book(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Please type in the genre you are interested in.')
    return ASK_GENRE

async def handle_genre(update: Update, context: ContextTypes.DEFAULT_TYPE):
    genre = update.message.text
    books = fetch_books_by_genre(genre)
    if not books:
        await update.message.reply_text(f"No books found for the genre: {genre}. Please try another genre.")
        return ConversationHandler.END
    
    csv_file = create_csv_file(books)
    await context.bot.send_document(chat_id=update.message.chat_id, document=csv_file, filename=f"{genre}_books.csv")
    return ConversationHandler.END

def fetch_books_by_genre(genre):
    url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={API_KEY}"
    response = requests.get(url)
    if response.status_code != 200:
        return None
    
    data = response.json()
    books = []
    for item in data.get('items', []):
        volume_info = item.get('volumeInfo', {})
        book = {
            "title": volume_info.get("title", "N/A"),
            "author": ", ".join(volume_info.get("authors", ["Unknown"])),
            "description": volume_info.get("description", "No description available."),
            "published_year": volume_info.get("publishedDate", "N/A"),
            "language": volume_info.get("language", "N/A"),
            "preview_link": volume_info.get("previewLink", "N/A")
        }
        books.append(book)
    return books

def create_csv_file(books):
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=["title", "author", "description", "published_year", "language", "preview_link"])
    writer.writeheader()
    writer.writerows(books)
    output.seek(0)
    return io.BytesIO(output.getvalue().encode())

async def preview(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Please provide the name of the book you need a preview link for.')
    return ASK_BOOK_NAME

async def handle_book_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    book_name = update.message.text
    preview_link = fetch_preview_link(book_name)
    
    if preview_link:
        message = f"Here's the preview link for <b>{book_name}</b>:\n<a href='{preview_link}'>{preview_link}</a>"
        await update.message.reply_text(message, parse_mode='HTML')
    else:
        await update.message.reply_text(f"No preview available for the book: {book_name}. Please try another book name.")
    
    return ConversationHandler.END

def fetch_preview_link(book_name):
    url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={API_KEY}"
    response = requests.get(url)
    if response.status_code != 200:
        return None
    
    data = response.json()
    if 'items' in data and len(data['items']) > 0:
        volume_info = data['items'][0].get('volumeInfo', {})
        return volume_info.get('previewLink', "No preview link available.")
    
    return None

async def list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Please type in the specific book name. You can use the /reading_list command to manage your list.')

async def reading_list(update: Update, context: ContextTypes.DEFAULT_TYPE):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add_book')],
        [InlineKeyboardButton("Delete a book", callback_data='delete_book')],
        [InlineKeyboardButton("View Reading List", callback_data='view_list')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Choose an action:', reply_markup=reply_markup)

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    help_text = (
        "/start -  Welcome to the Book Bot!\n"
        "/book -  Enter a genre to discover great books\n"
        "/preview -  Type a book name to get a preview link\n"
        "/list -  Manage your reading list with book names\n"
        "/reading_list -  View or modify your personalized reading list\n"
        "/help -  Get a list of all available commands"
    )
    await update.message.reply_text(help_text)

def main():
    initialize_reading_list()  

    application = Application.builder().token(BOT_TOKEN).build()

    book_handler = ConversationHandler(
        entry_points=[CommandHandler('book', book)],
        states={
            ASK_GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_genre)],
        },
        fallbacks=[]
    )

    preview_handler = ConversationHandler(
        entry_points=[CommandHandler('preview', preview)],
        states={
            ASK_BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_name)],
        },
        fallbacks=[]
    )

    application.add_handler(CommandHandler("start", start))
    application.add_handler(book_handler)
    application.add_handler(preview_handler)
    application.add_handler(CommandHandler("reading_list", reading_list))
    application.add_handler(CallbackQueryHandler(button))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(CommandHandler("list", list_command))
    application.add_handler(CommandHandler("help", help_command))

    print("Bot is running...")
    application.run_polling()

if __name__ == "__main__":
    main()
